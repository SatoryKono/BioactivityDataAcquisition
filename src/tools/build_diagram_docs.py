"""Build 5 Word documents with architecture diagrams grouped by layer.

Usage:
    python src/tools/build_diagram_docs.py
    python src/tools/build_diagram_docs.py --out reports/diagrams
    python src/tools/build_diagram_docs.py --only domain application
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

REPO_ROOT = Path(__file__).resolve().parents[2]
PNG_ROOT = REPO_ROOT / "docs/02-architecture/mmd-diagrams"

F = PNG_ROOT / "foundation/png"
A = PNG_ROOT / "architecture/png"
C = PNG_ROOT / "class-diagrams/png"

# ---------------------------------------------------------------------------
# Layer colours (RGB) — matches mermaid theme
# ---------------------------------------------------------------------------

LAYER_COLORS = {
    "domain": RGBColor(0x6A, 0x1B, 0x9A),  # purple
    "application": RGBColor(0x2E, 0x7D, 0x32),  # green
    "infrastructure": RGBColor(0xC6, 0x28, 0x28),  # red
    "composition": RGBColor(0xE6, 0x51, 0x00),  # orange
    "overview": RGBColor(0x15, 0x65, 0xC0),  # blue
}

# ---------------------------------------------------------------------------
# Manifest: 5 files × N sections × M diagrams
#
# Each entry:  (caption, png_path)
# Sections group diagrams visually with a heading.
# ---------------------------------------------------------------------------

MANIFEST: dict[str, dict] = {
    # ── 1. DOMAIN ────────────────────────────────────────────────────────────
    "domain": {
        "title": "BioETL Architecture — Domain Layer",
        "subtitle": "Contracts · Entities · Types · Exceptions",
        "sections": [
            (
                "Port Contracts",
                [
                    (
                        "Domain Ports — all 29 Protocol interfaces",
                        C / "01-domain-ports.png",
                    ),
                    (
                        "Port → Protocol Contracts (architecture view)",
                        A / "13-port-protocol-contracts.png",
                    ),
                    (
                        "Port → Adapter mapping reference",
                        F / "30-port-adapter-mapping.png",
                    ),
                    (
                        "Hexagonal Ports & Adapters",
                        F / "26-hexagonal-ports-adapters.png",
                    ),
                ],
            ),
            (
                "Class Diagrams",
                [
                    ("Entities & Aggregates", C / "02-entities-aggregates.png"),
                    ("Value Objects", C / "03-value-objects.png"),
                    ("Types & Enums", C / "04-types-enums.png"),
                    ("Configuration Classes", C / "06-config-classes.png"),
                    ("Domain Services", C / "13-domain-services.png"),
                ],
            ),
            (
                "Structure & DDD",
                [
                    (
                        "Domain Layer — ports, entities, config",
                        F / "04-domain-layer-class-diagram.png",
                    ),
                    ("Domain-Driven Design diagram", F / "08-domain-ddd.png"),
                    (
                        "Domain Model Relationships",
                        F / "13-domain-models-relationship.png",
                    ),
                ],
            ),
            (
                "Exceptions",
                [
                    (
                        "Exception Hierarchy (full tree)",
                        F / "50-exception-hierarchy.png",
                    ),
                    (
                        "Error Classification Tree",
                        F / "41-error-classification-tree.png",
                    ),
                    ("Exceptions class diagram", C / "05-exceptions.png"),
                ],
            ),
        ],
    },
    # ── 2. APPLICATION ───────────────────────────────────────────────────────
    "application": {
        "title": "BioETL Architecture — Application Layer",
        "subtitle": "Pipeline · Services · Transformers · Batch Execution",
        "sections": [
            (
                "Core Classes",
                [
                    (
                        "Application Core Services (PipelineRunner, BatchExecutor, LockManager)",
                        C / "07-application-core-services.png",
                    ),
                    (
                        "Application Services (DQ, Health, Export, Vacuum, Quarantine)",
                        C / "08-application-services.png",
                    ),
                    (
                        "PipelineRunner Class — all 14 DI dependencies",
                        F / "42-pipeline-runner-class.png",
                    ),
                    (
                        "Application Core Collaboration",
                        F / "40-application-core-collaboration.png",
                    ),
                ],
            ),
            (
                "Pipeline Lifecycle",
                [
                    (
                        "Pipeline Execution Flow (sequence)",
                        A / "04-pipeline-execution-flow.png",
                    ),
                    (
                        "Pipeline Execution — happy path",
                        F / "03-pipeline-execution-happy-path.png",
                    ),
                    ("Pipeline Sequence Diagram", F / "03-pipeline-sequence.png"),
                    (
                        "Pipeline Lifecycle States (FSM)",
                        F / "05-pipeline-lifecycle-states.png",
                    ),
                    ("PipelineRun Aggregate FSM", F / "31-pipeline-run-lifecycle.png"),
                    ("Complete ETL Workflow", F / "08-complete-etl-workflow.png"),
                ],
            ),
            (
                "Batch & Transform",
                [
                    ("BatchExecutor Internals", A / "15-batch-executor-internals.png"),
                    ("Transformer Hierarchy", A / "16-transformer-hierarchy.png"),
                    ("Transformers class diagram", C / "09-transformers.png"),
                    ("Extractors class diagram", C / "15-extractors.png"),
                    (
                        "Template Method — BaseTransformer hierarchy",
                        F / "45-template-method-transformer.png",
                    ),
                    ("Batch Processing Flow", F / "34-batch-processing-flow.png"),
                    (
                        "Fan-Out / Fan-In — asyncio.gather",
                        F / "43-fan-out-fan-in-pattern.png",
                    ),
                ],
            ),
            (
                "Data Quality",
                [
                    ("Data Quality System", A / "07-dq-system.png"),
                    ("DQ Check Workflow", F / "15-dq-check-workflow.png"),
                    ("Quarantine Record States", F / "20-quarantine-record-states.png"),
                ],
            ),
            (
                "Data Journey",
                [
                    (
                        "Single Record Journey — API → Bronze → Silver → Gold",
                        F / "32-single-record-journey.png",
                    ),
                    (
                        "Activity Entity Data Flow",
                        F / "21-activity-entity-data-flow.png",
                    ),
                    ("Pipeline Hierarchy", F / "17-pipeline-hierarchy.png"),
                ],
            ),
        ],
    },
    # ── 3. INFRASTRUCTURE ────────────────────────────────────────────────────
    "infrastructure": {
        "title": "BioETL Architecture — Infrastructure Layer",
        "subtitle": "Adapters · Storage · Resilience · Observability",
        "sections": [
            (
                "Adapters",
                [
                    ("Adapters class diagram", C / "10-adapters.png"),
                    (
                        "Provider Adapter Hierarchy",
                        A / "05-provider-adapter-hierarchy.png",
                    ),
                    (
                        "Infrastructure Layer Classes",
                        F / "10-infrastructure-layer-class-diagram.png",
                    ),
                    (
                        "Cross-Provider Enrichment (5-provider flow)",
                        F / "44-cross-provider-enrichment.png",
                    ),
                    (
                        "Client API Request Sequence",
                        F / "22-client-api-request-sequence.png",
                    ),
                ],
            ),
            (
                "Storage",
                [
                    (
                        "Storage Layer (Bronze/Silver/Gold/Delta)",
                        A / "06-storage-layer.png",
                    ),
                    ("Storage class diagram", C / "11-storage.png"),
                    ("Bronze Write Sequence", F / "18-bronze-write-sequence.png"),
                    (
                        "Delta Lake Write Sequence",
                        F / "19-delta-lake-write-sequence.png",
                    ),
                    ("SilverWriter Class", F / "23-silver-writer-class.png"),
                ],
            ),
            (
                "Resilience",
                [
                    (
                        "Resilience Patterns (circuit breaker, retry, rate limiter)",
                        A / "10-resilience-patterns.png",
                    ),
                    ("Circuit Breaker States", F / "07-circuit-breaker-states.png"),
                    (
                        "CircuitBreaker & Observer Class",
                        F / "25-circuit-breaker-observer-class.png",
                    ),
                    ("Provider Health States", F / "14-provider-health-states.png"),
                ],
            ),
            (
                "Security & Observability",
                [
                    ("Security, PII & Audit", A / "17-security-pii-audit.png"),
                    ("Hash Service Class", F / "24-hash-service-class.png"),
                    (
                        "Observability Stack (logging, metrics, tracing)",
                        A / "09-observability-stack.png",
                    ),
                    ("Observability class diagram", C / "14-observability.png"),
                ],
            ),
        ],
    },
    # ── 4. COMPOSITION ───────────────────────────────────────────────────────
    "composition": {
        "title": "BioETL Architecture — Composition Layer",
        "subtitle": "Bootstrap · DI · Configuration · Composite Pipeline",
        "sections": [
            (
                "Bootstrap & Dependency Injection",
                [
                    (
                        "Factories & Bootstrap class diagram",
                        C / "16-factories-bootstrap.png",
                    ),
                    ("Bootstrap / DI Container", A / "12-bootstrap-di-container.png"),
                    (
                        "Composition Root DI Graph",
                        F / "28-composition-root-di-graph.png",
                    ),
                    ("Bootstrap 9-step Sequence", F / "35-bootstrap-sequence.png"),
                    (
                        "Runtime Assembly Sequence (phases 1–8)",
                        F / "38-runtime-assembly-sequence.png",
                    ),
                ],
            ),
            (
                "Configuration",
                [
                    ("Configuration System", A / "11-configuration-system.png"),
                    (
                        "YAML Config Resolution — hierarchical merge",
                        F / "46-yaml-config-resolution.png",
                    ),
                ],
            ),
            (
                "Composite Pipeline",
                [
                    (
                        "Composite Pipeline class diagram",
                        C / "12-composite-pipeline.png",
                    ),
                    (
                        "Composite Pipeline (architecture)",
                        A / "08-composite-pipeline.png",
                    ),
                    (
                        "Composite Pipeline Workflow (Seed→Deps→FanOut→Merge→Gold)",
                        F / "29-composite-pipeline-workflow.png",
                    ),
                    (
                        "Composite Pipeline FSM — 10-state lifecycle",
                        F / "48-composite-phase-lifecycle.png",
                    ),
                    (
                        "CompositePipelineRunner component diagram",
                        F / "49-composite-runner-class.png",
                    ),
                    (
                        "Publication Composite — multi-source merge",
                        F / "47-publication-merge-sources.png",
                    ),
                ],
            ),
            (
                "Locking & Checkpoints",
                [
                    (
                        "Lock, Checkpoint & Shutdown",
                        A / "18-lock-checkpoint-shutdown.png",
                    ),
                    ("Locking Mechanism", F / "05-locking.png"),
                    (
                        "Lock Acquisition Sequence",
                        F / "11-lock-acquisition-sequence.png",
                    ),
                    ("MemoryLock Class", F / "16-memory-lock-class.png"),
                ],
            ),
        ],
    },
    # ── 5. OVERVIEW ──────────────────────────────────────────────────────────
    "overview": {
        "title": "BioETL Architecture — System Overview",
        "subtitle": "Hexagonal Architecture · Medallion · CLI · ER · Deployment",
        "sections": [
            (
                "System Overview",
                [
                    (
                        "High-Level Hexagonal Architecture",
                        A / "01-high-level-hexagonal.png",
                    ),
                    (
                        "Full System Component Diagram (C4-style)",
                        F / "01-full-system-component.png",
                    ),
                    ("High-Level System Overview", F / "01-high-level.png"),
                    (
                        "Architecture Principles Mindmap",
                        F / "36-architecture-principles-mindmap.png",
                    ),
                ],
            ),
            (
                "Architecture Rules",
                [
                    (
                        "Layer Dependency Matrix (ARCH-001)",
                        A / "02-layer-dependency-matrix.png",
                    ),
                    (
                        "Import Matrix Enforcement",
                        F / "27-import-matrix-enforcement.png",
                    ),
                    ("Layers Interaction", F / "05-layers-interaction.png"),
                ],
            ),
            (
                "Medallion Architecture",
                [
                    (
                        "Medallion Data Flow (Bronze→Silver→Gold)",
                        A / "03-medallion-data-flow.png",
                    ),
                    (
                        "Full Medallion Data Flow (detailed)",
                        F / "02-full-medallion-data-flow.png",
                    ),
                    ("Medallion (simplified)", F / "02-medallion.png"),
                    (
                        "Medallion Invariants — ARCH-007 RunType policy",
                        F / "39-medallion-invariants.png",
                    ),
                ],
            ),
            (
                "CLI / Interfaces",
                [
                    ("CLI Interface Layer", A / "14-cli-interface-layer.png"),
                    (
                        "CLI → PipelineRunnerService interaction",
                        F / "33-cli-run-interaction.png",
                    ),
                    (
                        "CLI Entry → Exit Code full chain",
                        F / "37-cli-entry-full-chain.png",
                    ),
                ],
            ),
            (
                "Data Model & Deployment",
                [
                    ("Full Entity-Relationship Diagram", F / "09-full-er-diagram.png"),
                    (
                        "Local Deployment Architecture (ADR-010)",
                        F / "12-local-deployment-architecture.png",
                    ),
                ],
            ),
        ],
    },
}

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type="page")  # type: ignore[arg-type]
    # Fallback via XML if above fails
    from docx.oxml.ns import qn as _qn

    br = OxmlElement("w:br")
    br.set(_qn("w:type"), "page")
    para.runs[0]._r.append(br) if para.runs else para._p.append(br)


def _set_heading_color(heading, color: RGBColor) -> None:
    for run in heading.runs:
        run.font.color.rgb = color


def _add_cover(doc: Document, title: str, subtitle: str, color: RGBColor) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(h, color)

    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph("BioETL Project · Architecture Documentation")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.size = Pt(10)
    info.runs[0].font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)


def _add_diagram(doc: Document, caption: str, png: Path, color: RGBColor) -> None:
    if not png.exists():
        p = doc.add_paragraph(f"[MISSING] {png.name}")
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        return

    # Caption above image
    cap = doc.add_paragraph(caption)
    cap.runs[0].bold = True
    cap.runs[0].font.size = Pt(11)
    cap.runs[0].font.color.rgb = color

    # Image — fit to page width
    doc.add_picture(str(png), width=Inches(6.3))

    # Filename note
    note = doc.add_paragraph(png.name)
    note.runs[0].font.size = Pt(8)
    note.runs[0].font.color.rgb = RGBColor(0xB0, 0xB0, 0xB0)
    note.runs[0].italic = True

    doc.add_paragraph()  # spacing


# ---------------------------------------------------------------------------
# Builder
# ---------------------------------------------------------------------------


def build_doc(layer: str, spec: dict, out_dir: Path) -> Path:
    color = LAYER_COLORS[layer]
    doc = Document()

    # Page margins — narrow for max diagram width
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    _add_cover(doc, spec["title"], spec["subtitle"], color)

    missing: list[str] = []

    for sec_idx, (sec_title, diagrams) in enumerate(spec["sections"]):
        if sec_idx > 0:
            pass  # no forced page break between sections

        # Section heading
        h2 = doc.add_heading(sec_title, level=1)
        _set_heading_color(h2, color)
        doc.add_paragraph()

        for caption, png in diagrams:
            if not png.exists():
                missing.append(str(png))
            _add_diagram(doc, caption, png, color)

    out_path = out_dir / f"{layer.upper()}.docx"
    doc.save(str(out_path))

    total = sum(len(d) for _, d in spec["sections"])
    found = total - len(missing)
    print(f"  [{layer.upper():15s}] {found:2d}/{total} diagrams -> {out_path.name}")
    if missing:
        for m in missing:
            print(f"    MISSING: {Path(m).name}")

    return out_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Build architecture Word docs from PNG diagrams"
    )
    parser.add_argument(
        "--out",
        default="reports/diagrams",
        help="Output directory (default: reports/diagrams)",
    )
    parser.add_argument(
        "--only",
        nargs="+",
        choices=list(MANIFEST.keys()),
        help="Build only specified layers",
    )
    args = parser.parse_args()

    out_dir = REPO_ROOT / args.out
    out_dir.mkdir(parents=True, exist_ok=True)

    layers = args.only or list(MANIFEST.keys())

    print("=" * 60)
    print("BUILDING ARCHITECTURE DIAGRAM DOCS")
    print(f"Output: {out_dir}")
    print("=" * 60)

    built: list[Path] = []
    for layer in layers:
        path = build_doc(layer, MANIFEST[layer], out_dir)
        built.append(path)

    print("=" * 60)
    print(f"Done. {len(built)} file(s) written to {out_dir}")


if __name__ == "__main__":
    main()
